import os
import json

from django.http import JsonResponse, HttpResponse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods

from requests import post

from docx import Document
from api.ocr import process_image 
from api.restore import enhance_to_bytes, enchance_and_save, HSV_ADJUSTMENT
from api import s3

AUTO_ENCHANCE = True

@csrf_exempt
def healthcheck(request):
    return JsonResponse({"status": "ok"})


@require_http_methods(["POST"])
@csrf_exempt
def upload(request):
    uploaded_file = request.FILES['file']

    if AUTO_ENCHANCE:
        image = enhance_to_bytes(uploaded_file)
    else:
        image = uploaded_file.read()

    try:
        response_data = process_image(image)
        print(response_data)
        return JsonResponse(response_data)
    except Exception as e:
        return JsonResponse({'status': 'error', 'message': str(e)}, status=500)


@require_http_methods(["POST"])
@csrf_exempt
def restore(request):
    key = s3.generate_key('image', 'jpg', 'file.jpg')
    path = "media/"+key
    enchance_and_save(request.FILES['file'], path)
    file_url = s3.upload_file(path, key)
    return JsonResponse({"file_url": file_url})


def extract_fields(data):
    ocr = data.get('ocr')
    enchanced_text = data.get('enhancedText')

    analysis = data.get('analysis')
    attributes = analysis.get('attributes')
    attributes_names = " ".join(attributes.get('names'))
    attributes_dates = " ".join(attributes.get('dates'))
    attributes_places = " ".join(attributes.get('places'))

    structured_content = data.get('structuredContent')
    summary = structured_content.get('summary')
    details = structured_content.get('details')
    details_names = " ".join(details.get('names'))
    details_dates = " ".join(details.get('dates'))
    details_places = " ".join(details.get('places'))

    return {
        "OCR": ocr,
        "Enchanced text": enchanced_text,
        "Attributes/names": attributes_names,
        "Attributes/dates": attributes_dates,
        "Attributes/places": attributes_places,
        "Summary": summary,
        "Details/names": details_names,
        "Details/dates": details_dates,
        "Details/places": details_places,
    }


def docx(data):
    fields = extract_fields(data)
    document = Document()
    for key, value in fields.items():
        if not value:
            continue
        document.add_heading(key, level=1)
        document.add_paragraph(value)

    key = s3.generate_key('author', 'docx', 'file.docx')
    document.save("media/"+key)
    return key


@require_http_methods(["POST"])
@csrf_exempt
def to_docx(request):
    data = json.loads(request.body)
    key = docx(data)
    file_url = s3.upload_file("media/"+key, key)
    return JsonResponse({"file_url": file_url})


@require_http_methods(["GET"])
@csrf_exempt
def swap_auto_enchance(request):
    global AUTO_ENCHANCE
    AUTO_ENCHANCE = not AUTO_ENCHANCE
    return JsonResponse({"enchance_mode": AUTO_ENCHANCE})


@require_http_methods(["GET"])
@csrf_exempt
def swap_hsv_adjustment(request):
    global HSV_ADJUSTMENT
    HSV_ADJUSTMENT = not HSV_ADJUSTMENT
    return JsonResponse({"hsv_adjustment": HSV_ADJUSTMENT})
